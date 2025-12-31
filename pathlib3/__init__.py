#!/usr/bin/env python3

# File: pathlib3/__init__.py
# Author: Hadi Cahyadi <cumulus13@gmail.com>
# Date: 2025-12-30
# Description: Extended pathlib with additional utility methods
# License: MIT

"""
pathlib3 - Extended pathlib with additional utility methods

This module re-exports everything from pathlib and adds Path class
with extended functionality.

Usage:
    from pathlib3 import Path
    # or
    from pathlib3 import *
"""

# ===================================================================
# IMPORT EVERYTHING from pathlib
# ===================================================================
from pathlib import (
    # Main classes
    Path as _PathBase,
    PurePath,
    PosixPath,
    WindowsPath,
    PurePosixPath,
    PureWindowsPath,
)

# Import modules for extended functionality
import pathlib
import os
import sys
import shutil
import hashlib
import json
import pickle
from typing import Union, Tuple, List, Optional, Callable, Any, Iterator
from datetime import datetime
import traceback

RICH_AVAILABLE = False
MUTAGEN_AVAILABLE = False
YAML_AVAILABLE = False
TOML_AVAILABLE = False
INI_AVAILABLE = False
PIL_AVAILABLE = False
PYPDF2_AVAILABLE = False
PYTHON_DOCX_AVAILABLE = False
OPENPYXL_AVAILABLE = False
EMAIL_AVAILABLE = False

# Check for optional dependencies
try:
    from rich.table import Table
    from rich.console import Console
    console = Console(width=os.get_terminal_size().columns)
    RICH_AVAILABLE = True
except:
    Table = None

try:
    from mutagen import File as MutagenFile  # type: ignore
    from mutagen.id3 import ID3, APIC, error as ID3Error  # type: ignore
    MUTAGEN_AVAILABLE = True
except:
    pass

try:
    import yaml
    YAML_AVAILABLE = True
except ImportError:
    pass

try:
    if sys.version_info < (3, 11):
        import tomli 
    else:
        import tomllib
    TOML_AVAILABLE = True
except ImportError:
    try:
        import tomllib
        TOML_AVAILABLE = True
    except ImportError:
        pass

try:
    import configparser
    INI_AVAILABLE = True  # Built-in, always available
except ImportError:
    pass

try:
    from PIL import Image
    from PIL.ExifTags import TAGS
    PIL_AVAILABLE = True
except ImportError:
    pass

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    pass

try:
    import docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    pass

try:
    from openpyxl import load_workbook
    OPENPYXL_AVAILABLE = True
except ImportError:
    pass

# Check for email functionality
try:
    import smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart
    from email.mime.base import MIMEBase
    from email import encoders
    from email.mime.image import MIMEImage
    from email.mime.audio import MIMEAudio
    from email.mime.application import MIMEApplication
    EMAIL_AVAILABLE = True  # Built-in, always available
except ImportError:
    pass

# ===================================================================
# Email Configuration Helper
# ===================================================================
class EmailConfig:
    """
    Email configuration helper for SMTP settings.
    
    Common SMTP servers:
        Gmail: smtp.gmail.com:587 (TLS) or :465 (SSL)
        Outlook: smtp-mail.outlook.com:587
        Yahoo: smtp.mail.yahoo.com:587
        Office365: smtp.office365.com:587
    
    Example:
        >>> config = EmailConfig(
        ...     smtp_server='smtp.gmail.com',
        ...     smtp_port=587,
        ...     username='your.email@gmail.com',
        ...     password='your_app_password',
        ...     use_tls=True
        ... )
    """
    
    def __init__(
        self,
        smtp_server: str,
        smtp_port: int = 587,
        username: str = '',
        password: str = '',
        use_tls: bool = True,
        use_ssl: bool = False,
        timeout: int = 30
    ):
        """
        Initialize email configuration.
        
        Args:
            smtp_server: SMTP server address
            smtp_port: SMTP port (587 for TLS, 465 for SSL, 25 for no encryption)
            username: Email username/address
            password: Email password or app-specific password
            use_tls: Use TLS encryption (STARTTLS)
            use_ssl: Use SSL encryption
            timeout: Connection timeout in seconds
        """
        self.smtp_server = smtp_server
        self.smtp_port = smtp_port
        self.username = username
        self.password = password
        self.use_tls = use_tls
        self.use_ssl = use_ssl
        self.timeout = timeout
    
    @classmethod
    def gmail(cls, username: str, password: str) -> 'EmailConfig':
        """Quick config for Gmail"""
        return cls(
            smtp_server='smtp.gmail.com',
            smtp_port=587,
            username=username,
            password=password,
            use_tls=True
        )
    
    @classmethod
    def outlook(cls, username: str, password: str) -> 'EmailConfig':
        """Quick config for Outlook/Hotmail"""
        return cls(
            smtp_server='smtp-mail.outlook.com',
            smtp_port=587,
            username=username,
            password=password,
            use_tls=True
        )
    
    @classmethod
    def office365(cls, username: str, password: str) -> 'EmailConfig':
        """Quick config for Office 365"""
        return cls(
            smtp_server='smtp.office365.com',
            smtp_port=587,
            username=username,
            password=password,
            use_tls=True
        )
    
    @classmethod
    def yahoo(cls, username: str, password: str) -> 'EmailConfig':
        """Quick config for Yahoo"""
        return cls(
            smtp_server='smtp.mail.yahoo.com',
            smtp_port=587,
            username=username,
            password=password,
            use_tls=True
        )


# ===================================================================
# Path - Extended Path Class
# ===================================================================
class Path(type(_PathBase())):
    """
    Extended Path class with 40+ additional utility methods.
    Inherits ALL functionality from pathlib.Path and adds more.
    
    All original pathlib.Path methods are available:
    - .exists(), .is_file(), .is_dir(), .is_symlink()
    - .stat(), .chmod(), .rename(), .replace()
    - .mkdir(), .rmdir(), .unlink(), .touch()
    - .read_text(), .read_bytes(), .write_text(), .write_bytes()
    - .glob(), .rglob(), .iterdir()
    - .resolve(), .absolute(), .relative_to()
    - .with_name(), .with_suffix(), .with_stem()
    - .parent, .parents, .name, .stem, .suffix, .suffixes
    - .anchor, .parts, .drive, .root
    - .as_posix(), .as_uri()
    - .expanduser(), .home(), .cwd()
    - .match(), .is_relative_to(), .is_absolute()
    - .joinpath(), .samefile()
    - and many more...
    
    Additional method categories:
    
    BASIC UTILITIES:
    - .ext() - Get extension without dot
    - .basename() - Get filename with extension
    - .base() - Get filename without extension
    - .dirname() - Get directory path as string
    - .abspath() - Get absolute path as string

    NONE HANDLING:
    - Path(None) - Creates Path('.') instead of error
    - .safe() - Class method to create Path safely from optional
    - .from_optional() - Class method to create Path or return None
    
    PATH MANIPULATION:
    - .normpath() - Normalize path
    - .join() - Join path components
    - .split_ext() - Split into base and extension
    - .split_path() - Split path into components
    - .change_ext() - Change file extension
    
    DIRECTORY OPERATIONS:
    - .ensure_dir() - Create directory if doesn't exist
    - .ensure_parent() - Create parent directory
    - .touch_parent() - Create parent dirs and touch file
    - .ls() - List directory contents
    - .tree() - Show directory tree
    - .find() - Find files recursively
    
    FILE OPERATIONS:
    - .rm() - Remove file or directory
    - .copy_to() - Copy file to destination
    - .move_to() - Move file to destination
    - .append_text() - Append text to file
    - .append_bytes() - Append bytes to file
    - .backup() - Create backup of file
    
    FILE INFO:
    - .size() - Get file size
    - .size_human() - Get human-readable size
    - .mtime() - Get modification time
    - .ctime() - Get creation time
    - .atime() - Get access time
    - .age() - Get file age in seconds
    - .is_empty() - Check if empty
    - .is_newer_than() - Compare modification times
    - .is_older_than() - Compare modification times
    
    CONTENT OPERATIONS:
    - .lines() - Read lines as list
    - .read_json() - Read JSON file
    - .write_json() - Write JSON file
    - .read_pickle() - Read pickle file
    - .write_pickle() - Write pickle file
    - .hash() - Calculate file hash
    - .checksum() - Alias for hash
    - .count_lines() - Count lines in file
    - .validate() - Validate file format
    - .metadata() - Extract file metadata (images, PDFs, audio, video, docs)
    - .metadata_simple() - Get simple metadata summary
    
    SEARCH & FILTER:
    - .find_files() - Find files by pattern
    - .find_dirs() - Find directories by pattern
    - .walk() - Walk directory tree

    MUSIC TAGS:
    - .music_tag() - Get music tag
    - .show_info() - Display music tags
    
    COMPARISON:
    - .same_content() - Check if files have same content

    EMAIL OPERATIONS:
    - .email_as_attachment() - Send file as email attachment
    - .send_email() - Send email with multiple attachments (static method)

    """
    
    # ===============================================================
    # BASIC UTILITY METHODS
    # ===============================================================

    def __new__(cls, *args, **kwargs):
        """Create new Path instance with None handling."""
        # Handle None case
        if len(args) == 1 and args[0] is None:
            args = ('.',)
        elif len(args) > 1:
            # Filter out None values from args
            args = tuple(arg if arg is not None else '.' for arg in args)
        
        # Handle empty args
        if len(args) == 0:
            args = ('.',)
        
        return super().__new__(cls, *args, **kwargs)
    
    @classmethod
    def safe(cls, path: Optional[Union[str, _PathBase, 'Path']], default: str = '.') -> 'Path':
        """
        Create Path safely, handling None values.
        More explicit alternative to Path(None).
        
        Args:
            path: Path string or Path object (can be None)
            default: Default path if input is None (default: '.')
        
        Returns:
            Path: Path instance
        
        Example:
            >>> Path.safe(None)
            Path('.')
            >>> Path.safe(None, '/tmp')
            Path('/tmp')
            >>> Path.safe("file.txt")
            Path('file.txt')
        """
        if path is None:
            return cls(default)
        return cls(path)
    
    @classmethod
    def from_optional(cls, path: Optional[Union[str, _PathBase, 'Path']]) -> Optional['Path']:
        """
        Create Path from optional value, returns None if input is None.
        
        Args:
            path: Path string or Path object (can be None)
        
        Returns:
            Path or None: Path instance or None
        
        Example:
            >>> Path.from_optional(None)
            None
            >>> Path.from_optional("file.txt")
            Path('file.txt')
        """
        if path is None:
            return None
        return cls(path)
    
    def ext(self) -> str:
        """
        Get file extension without the dot.
        
        Returns:
            str: File extension (e.g., 'txt', 'py', 'tar.gz')
        
        Example:
            >>> Path('file.txt').ext()
            'txt'
            >>> Path('archive.tar.gz').ext()
            'gz'
        """
        return self.suffix.lstrip('.')
    
    def basename(self) -> str:
        """
        Get the base name (filename with extension).
        Alias for .name property.
        
        Returns:
            str: Basename (e.g., 'file.txt')
        
        Example:
            >>> Path('/home/user/file.txt').basename()
            'file.txt'
        """
        return self.name
    
    def base(self) -> str:
        """
        Get the base name without extension.
        Alias for .stem property.
        
        Returns:
            str: Filename without extension (e.g., 'file')
        
        Example:
            >>> Path('file.txt').base()
            'file'
        """
        return self.stem
    
    def dirname(self) -> str:
        """
        Get the directory name as string.
        
        Returns:
            str: Directory path
        
        Example:
            >>> Path('/home/user/file.txt').dirname()
            '/home/user'
        """
        return str(self.parent)
    
    def abspath(self) -> str:
        """
        Get absolute path as string.
        
        Returns:
            str: Absolute path
        
        Example:
            >>> Path('file.txt').abspath()
            '/current/working/dir/file.txt'
        """
        return str(self.absolute())
    
    # ===============================================================
    # PATH MANIPULATION
    # ===============================================================
    
    def normpath(self) -> 'Path':
        """
        Normalize path (remove redundant separators and up-level references).
        
        Returns:
            Path: Normalized path
        
        Example:
            >>> Path('/home//user/../user/file.txt').normpath()
            Path('/home/user/file.txt')
        """
        return Path(os.path.normpath(self))
    
    def join(self, *args) -> 'Path':
        """
        Join path components.
        
        Args:
            *args: Path components to join
        
        Returns:
            Path: Joined path
        
        Example:
            >>> Path('/home').join('user', 'documents', 'file.txt')
            Path('/home/user/documents/file.txt')
        """
        result = self
        for arg in args:
            result = result / arg
        return Path(result)
    
    def split_ext(self) -> Tuple[str, str]:
        """
        Split path into base and extension.
        
        Returns:
            tuple: (base_path, extension)
        
        Example:
            >>> Path('/home/user/file.txt').split_ext()
            ('/home/user/file', '.txt')
        """
        return (str(self.with_suffix('')), self.suffix)
    
    def split_path(self) -> List[str]:
        """
        Split path into list of components.
        
        Returns:
            list: List of path components
        
        Example:
            >>> Path('/home/user/file.txt').split_path()
            ['/', 'home', 'user', 'file.txt']
        """
        return list(self.parts)
    
    def change_ext(self, new_ext: str) -> 'Path':
        """
        Change file extension.
        
        Args:
            new_ext: New extension (with or without dot)
        
        Returns:
            Path: Path with new extension
        
        Example:
            >>> Path('file.txt').change_ext('md')
            Path('file.md')
            >>> Path('file.txt').change_ext('.json')
            Path('file.json')
        """
        if not new_ext.startswith('.'):
            new_ext = '.' + new_ext
        return Path(self.with_suffix(new_ext))
    
    # ===============================================================
    # DIRECTORY OPERATIONS
    # ===============================================================
    
    def ensure_dir(self, mode: int = 0o777) -> 'Path':
        """
        Create directory if it doesn't exist (for directories).
        
        Args:
            mode: Directory permissions (default: 0o777)
        
        Returns:
            Path: self (for chaining)
        
        Example:
            >>> Path('/tmp/new_folder').ensure_dir()
            Path('/tmp/new_folder')
        """
        self.mkdir(mode=mode, parents=True, exist_ok=True)
        return self
    
    def ensure_parent(self, mode: int = 0o777) -> 'Path':
        """
        Create parent directory if it doesn't exist.
        
        Args:
            mode: Directory permissions (default: 0o777)
        
        Returns:
            Path: self (for chaining)
        
        Example:
            >>> Path('/tmp/new_folder/file.txt').ensure_parent()
            Path('/tmp/new_folder/file.txt')
        """
        self.parent.mkdir(mode=mode, parents=True, exist_ok=True)
        return self
    
    def touch_parent(self) -> 'Path':
        """
        Create parent directories and touch file.
        
        Returns:
            Path: self (for chaining)
        
        Example:
            >>> Path('/tmp/new_folder/file.txt').touch_parent()
            Path('/tmp/new_folder/file.txt')
        """
        self.parent.mkdir(parents=True, exist_ok=True)
        self.touch()
        return self
    
    def ls(self, pattern: str = "*", only_files: bool = False, only_dirs: bool = False) -> List['Path']:
        """
        List directory contents as Path objects.
        
        Args:
            pattern: Glob pattern (default: "*")
            only_files: Only return files
            only_dirs: Only return directories
        
        Returns:
            list: List of Path objects
        
        Example:
            >>> Path('/tmp').ls()
            [Path('/tmp/file1.txt'), Path('/tmp/file2.txt')]
            >>> Path('/tmp').ls('*.txt')
            [Path('/tmp/file1.txt'), Path('/tmp/file2.txt')]
            >>> Path('/tmp').ls(only_files=True)
            [Path('/tmp/file1.txt')]
        """
        results = [Path(p) for p in self.glob(pattern)]
        
        if only_files:
            results = [p for p in results if p.is_file()]
        elif only_dirs:
            results = [p for p in results if p.is_dir()]
        
        return results
    
    def tree(self, max_depth: Optional[int] = None, prefix: str = "") -> str:
        """
        Generate directory tree structure as string.
        
        Args:
            max_depth: Maximum depth to traverse
            prefix: Prefix for formatting (internal use)
        
        Returns:
            str: Tree structure
        
        Example:
            >>> print(Path('/tmp').tree(max_depth=2))
            /tmp
            ├── file1.txt
            ├── folder1/
            │   ├── file2.txt
            │   └── file3.txt
        """
        if not self.is_dir():
            return str(self)
        
        lines = [str(self)]
        
        if max_depth is not None and max_depth <= 0:
            return lines[0]
        
        try:
            entries = sorted(self.iterdir(), key=lambda x: (not x.is_dir(), x.name))
            for i, entry in enumerate(entries):
                is_last = i == len(entries) - 1
                current_prefix = "└── " if is_last else "├── "
                child_prefix = "    " if is_last else "│   "
                
                if entry.is_dir():
                    lines.append(f"{prefix}{current_prefix}{entry.name}/")
                    if max_depth is None or max_depth > 1:
                        child_tree = Path(entry).tree(
                            max_depth=max_depth - 1 if max_depth else None,
                            prefix=prefix + child_prefix
                        )
                        child_lines = child_tree.split('\n')[1:]  # Skip first line
                        lines.extend(child_lines)
                else:
                    lines.append(f"{prefix}{current_prefix}{entry.name}")
        except PermissionError:
            lines.append(f"{prefix}[Permission Denied]")
        
        return '\n'.join(lines)
    
    def find(self, pattern: str = "*", recursive: bool = True) -> List['Path']:
        """
        Find files matching pattern.
        
        Args:
            pattern: Glob pattern
            recursive: Search recursively
        
        Returns:
            list: List of matching Path objects
        
        Example:
            >>> Path('/tmp').find('*.txt')
            [Path('/tmp/file1.txt'), Path('/tmp/sub/file2.txt')]
        """
        if recursive:
            return [Path(p) for p in self.rglob(pattern)]
        return [Path(p) for p in self.glob(pattern)]
    
    # ===============================================================
    # FILE OPERATIONS
    # ===============================================================
    
    def rm(self, recursive: bool = False, missing_ok: bool = False) -> None:
        """
        Remove file or directory.
        
        Args:
            recursive: If True, remove directory recursively
            missing_ok: If True, don't raise error if path doesn't exist
        
        Example:
            >>> Path('file.txt').rm()
            >>> Path('folder').rm(recursive=True)
        """
        if not self.exists():
            if missing_ok:
                return
            raise FileNotFoundError(f"{self} does not exist")
        
        if self.is_file() or self.is_symlink():
            self.unlink()
        elif self.is_dir():
            if recursive:
                shutil.rmtree(self)
            else:
                self.rmdir()
    
    def copy_to(self, dest: Union[str, _PathBase, 'Path'], overwrite: bool = False) -> 'Path':
        """
        Copy file or directory to destination.
        
        Args:
            dest: Destination path
            overwrite: Overwrite if exists
        
        Returns:
            Path: Destination path
        
        Example:
            >>> Path('source.txt').copy_to('dest.txt')
            Path('dest.txt')
        """
        dest = Path(dest)
        
        if dest.exists() and not overwrite:
            raise FileExistsError(f"{dest} already exists")
        
        if self.is_file():
            dest.ensure_parent()
            shutil.copy2(self, dest)
        elif self.is_dir():
            if dest.exists():
                shutil.rmtree(dest)
            shutil.copytree(self, dest)
        
        return dest
    
    def move_to(self, dest: Union[str, _PathBase, 'Path']) -> 'Path':
        """
        Move file or directory to destination.
        
        Args:
            dest: Destination path
        
        Returns:
            Path: Destination path
        
        Example:
            >>> Path('old.txt').move_to('new.txt')
            Path('new.txt')
        """
        dest = Path(dest)
        dest.ensure_parent()
        shutil.move(str(self), str(dest))
        return dest
    
    def append_text(self, text: str, encoding: str = 'utf-8', newline: bool = False) -> 'Path':
        """
        Append text to file.
        
        Args:
            text: Text to append
            encoding: Text encoding (default: 'utf-8')
            newline: Add newline before text
        
        Returns:
            Path: self (for chaining)
        
        Example:
            >>> Path('log.txt').append_text('New log entry')
        """
        mode = 'a'
        try:
            with self.open(mode, encoding=encoding) as f:
                if newline and self.exists() and self.stat().st_size > 0:
                    f.write('\n')
                f.write(text)
        except PermissionError:
            raise PermissionError(f"Permission denied: '{self}'. File may be open in another program.")
        return self
    
    def append_bytes(self, data: bytes) -> 'Path':
        """
        Append bytes to file.
        
        Args:
            data: Bytes to append
        
        Returns:
            Path: self (for chaining)
        
        Example:
            >>> Path('data.bin').append_bytes(b'\\x00\\x01\\x02')
        """
        try:
            with self.open('ab') as f:
                f.write(data)
        except PermissionError:
            raise PermissionError(f"Permission denied: '{self}'. File may be open in another program.")
        return self
    
    def backup(self, suffix: str = '.bak') -> 'Path':
        """
        Create backup of file.
        
        Args:
            suffix: Backup suffix (default: '.bak')
        
        Returns:
            Path: Backup file path
        
        Example:
            >>> Path('important.txt').backup()
            Path('important.txt.bak')
        """
        backup_path = Path(str(self) + suffix)
        return self.copy_to(backup_path, overwrite=True)
    
    # ===============================================================
    # FILE INFOsd
    # ===============================================================
    
    def size(self) -> int:
        """
        Get file size in bytes. For directories, returns total size.
        
        Returns:
            int: Size in bytes
        
        Example:
            >>> Path('file.txt').size()
            1024
        """
        if not self.exists():
            return 0
        
        if self.is_file():
            return self.stat().st_size
        elif self.is_dir():
            total = 0
            try:
                for entry in self.rglob('*'):
                    if entry.is_file():
                        try:
                            total += entry.stat().st_size
                        except (OSError, PermissionError):
                            continue
            except (OSError, PermissionError):
                pass
            return total
        return 0
    
    def size_human(self) -> str:
        """
        Get human-readable file size.
        
        Returns:
            str: Human-readable size (e.g., '1.5 MB')
        
        Example:
            >>> Path('file.txt').size_human()
            '1.0 KB'
        """
        size = self.size()
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size < 1024.0:
                return f"{size:.1f} {unit}"
            size /= 1024.0
        return f"{size:.1f} PB"
    
    def mtime(self) -> float:
        """
        Get modification time as timestamp.
        
        Returns:
            float: Modification timestamp
        
        Example:
            >>> Path('file.txt').mtime()
            1704067200.0
        """
        return self.stat().st_mtime
    
    def ctime(self) -> float:
        """
        Get creation/metadata change time as timestamp.
        
        Returns:
            float: Creation timestamp
        
        Example:
            >>> Path('file.txt').ctime()
            1704067200.0
        """
        return self.stat().st_ctime
    
    def atime(self) -> float:
        """
        Get access time as timestamp.
        
        Returns:
            float: Access timestamp
        
        Example:
            >>> Path('file.txt').atime()
            1704067200.0
        """
        return self.stat().st_atime
    
    def age(self) -> float:
        """
        Get file age in seconds since last modification.
        
        Returns:
            float: Age in seconds
        
        Example:
            >>> Path('file.txt').age()
            3600.5
        """
        return datetime.now().timestamp() - self.mtime()
    
    def is_empty(self) -> bool:
        """
        Check if file or directory is empty.
        
        Returns:
            bool: True if empty, False otherwise
        
        Example:
            >>> Path('empty.txt').is_empty()
            True
        """
        if not self.exists():
            return True
        
        if self.is_file():
            return self.stat().st_size == 0
        elif self.is_dir():
            try:
                return not any(self.iterdir())
            except PermissionError:
                return False
        return False
    
    def is_newer_than(self, other: Union[str, _PathBase, 'Path']) -> bool:
        """
        Check if this file is newer than another.
        
        Args:
            other: Other file path
        
        Returns:
            bool: True if this file is newer
        
        Example:
            >>> Path('new.txt').is_newer_than('old.txt')
            True
        """
        other = Path(other)
        if not self.exists() or not other.exists():
            return False
        return self.mtime() > other.mtime()
    
    def is_older_than(self, other: Union[str, _PathBase, 'Path']) -> bool:
        """
        Check if this file is older than another.
        
        Args:
            other: Other file path
        
        Returns:
            bool: True if this file is older
        
        Example:
            >>> Path('old.txt').is_older_than('new.txt')
            True
        """
        other = Path(other)
        if not self.exists() or not other.exists():
            return False
        return self.mtime() < other.mtime()
    
    # ===============================================================
    # CONTENT OPERATIONS
    # ===============================================================
    
    def lines(self, encoding: str = 'utf-8', strip: bool = True, skip_empty: bool = False) -> List[str]:
        """
        Read file lines as list.
        
        Args:
            encoding: Text encoding (default: 'utf-8')
            strip: Strip whitespace from lines
            skip_empty: Skip empty lines
        
        Returns:
            list: List of lines
        
        Example:
            >>> Path('file.txt').lines()
            ['line 1', 'line 2', 'line 3']
        """
        try:
            lines = self.read_text(encoding=encoding).splitlines()
            if strip:
                lines = [line.strip() for line in lines]
            if skip_empty:
                lines = [line for line in lines if line]
            return lines
        except PermissionError:
            raise PermissionError(f"Permission denied: '{self}'. File may be open in another program.")
        except Exception as e:
            raise IOError(f"Error reading file '{self}': {e}")
    
    def read_json(self, encoding: str = 'utf-8', **kwargs) -> Any:
        """
        Read JSON file.
        
        Args:
            encoding: Text encoding (default: 'utf-8')
            **kwargs: Additional arguments for json.load()
        
        Returns:
            Parsed JSON data
        
        Example:
            >>> Path('config.json').read_json()
            {'key': 'value'}
        """
        try:
            return json.loads(self.read_text(encoding=encoding), **kwargs)
        except json.JSONDecodeError as e:
            raise ValueError(f"Invalid JSON in file '{self}': {e}")
        except PermissionError:
            raise PermissionError(f"Permission denied: '{self}'. File may be open in another program.")
    
    def write_json(self, data: Any, encoding: str = 'utf-8', indent: int = 2, ensure_ascii: bool = False, **kwargs) -> 'Path':
        """
        Write data as JSON file.
        
        Args:
            data: Data to write
            encoding: Text encoding (default: 'utf-8')
            indent: JSON indentation (default: 2)
            ensure_ascii: Ensure ASCII encoding (default: False)
            **kwargs: Additional arguments for json.dump()
        
        Returns:
            Path: self (for chaining)
        
        Example:
            >>> Path('config.json').write_json({'key': 'value'})
        """
        try:
            self.write_text(
                json.dumps(data, indent=indent, ensure_ascii=ensure_ascii, **kwargs), 
                encoding=encoding
            )
        except PermissionError:
            raise PermissionError(f"Permission denied: '{self}'. File may be open in another program.")
        return self
    
    def read_pickle(self) -> Any:
        """
        Read pickle file.
        
        Returns:
            Unpickled data
        
        Example:
            >>> Path('data.pkl').read_pickle()
            {'key': 'value'}
        """
        try:
            return pickle.loads(self.read_bytes())
        except PermissionError:
            raise PermissionError(f"Permission denied: '{self}'. File may be open in another program.")
        except pickle.UnpicklingError as e:
            raise ValueError(f"Invalid pickle data in file '{self}': {e}")
    
    def write_pickle(self, data: Any, protocol: int = pickle.HIGHEST_PROTOCOL) -> 'Path':
        """
        Write data as pickle file.
        
        Args:
            data: Data to pickle
            protocol: Pickle protocol version
        
        Returns:
            Path: self (for chaining)
        
        Example:
            >>> Path('data.pkl').write_pickle({'key': 'value'})
        """
        try:
            self.write_bytes(pickle.dumps(data, protocol=protocol))
        except PermissionError:
            raise PermissionError(f"Permission denied: '{self}'. File may be open in another program.")
        return self
    
    def hash(self, algorithm: str = 'sha256', chunk_size: int = 8192) -> str:
        """
        Calculate file hash.
        
        Args:
            algorithm: Hash algorithm (md5, sha1, sha256, etc.)
            chunk_size: Size of chunks to read (default: 8192 bytes)
        
        Returns:
            str: Hexadecimal hash
        
        Example:
            >>> Path('file.txt').hash()
            'a591a6d40bf420404a011733cfb7b190d62c65bf0bcda32b57b277d9ad9f146e'
        """
        if not self.is_file():
            raise ValueError(f"Cannot hash non-file: '{self}'")
        
        try:
            h = hashlib.new(algorithm)
            with self.open('rb') as f:
                for chunk in iter(lambda: f.read(chunk_size), b''):
                    h.update(chunk)
            return h.hexdigest()
        except PermissionError:
            raise PermissionError(f"Permission denied: '{self}'. File may be open in another program.")
        except ValueError as e:
            raise ValueError(f"Invalid hash algorithm '{algorithm}': {e}")
    
    def checksum(self, algorithm: str = 'sha256') -> str:
        """
        Alias for hash().
        
        Args:
            algorithm: Hash algorithm (default: 'sha256')
        
        Returns:
            str: Hexadecimal hash
        
        Example:
            >>> Path('file.txt').checksum()
            'a591a6d40bf420404a011733cfb7b190d62c65bf0bcda32b57b277d9ad9f146e'
        """
        return self.hash(algorithm)
    
    def count_lines(self, encoding: str = 'utf-8') -> int:
        """
        Count lines in file.
        
        Args:
            encoding: Text encoding (default: 'utf-8')
        
        Returns:
            int: Number of lines
        
        Example:
            >>> Path('file.txt').count_lines()
            42
        """
        if not self.is_file():
            raise ValueError(f"Cannot count lines in non-file: '{self}'")
        
        try:
            with self.open('r', encoding=encoding) as f:
                return sum(1 for _ in f)
        except PermissionError:
            raise PermissionError(f"Permission denied: '{self}'. File may be open in another program.")
        except UnicodeDecodeError as e:
            raise ValueError(f"Cannot decode file '{self}' with encoding '{encoding}': {e}")
    
    def validate(self, file_type: Optional[str] = None, strict: bool = True) -> Tuple[bool, Optional[str]]:
        """
        Validate file format (JSON, YAML, TOML, INI).
        
        Args:
            file_type: File type to validate ('json', 'yaml', 'yml', 'toml', 'ini')
                      If None, auto-detect from extension
            strict: If True, raise error for missing libraries. If False, return (False, error_msg)
        
        Returns:
            tuple: (is_valid, error_message)
                   (True, None) if valid
                   (False, "error message") if invalid
        
        Raises:
            ImportError: If required library is not installed (when strict=True)
            ValueError: If file type is not supported
        
        Example:
            >>> Path('config.json').validate()
            (True, None)
            >>> Path('config.yaml').validate()
            (False, "PyYAML library not installed. Install with: pip install pyyaml")
            >>> is_valid, error = Path('config.toml').validate(strict=False)
            >>> if not is_valid:
            ...     print(f"Invalid: {error}")
        
        Supported formats:
            - JSON (built-in, always available)
            - YAML (requires PyYAML: pip install pyyaml)
            - TOML (built-in Python 3.11+, or requires tomli: pip install tomli)
            - INI (built-in, always available)
        """
        # Auto-detect file type from extension
        if file_type is None:
            ext = self.ext().lower()
            if ext in ('yaml', 'yml'):
                file_type = 'yaml'
            elif ext == 'toml':
                file_type = 'toml'
            elif ext == 'ini':
                file_type = 'ini'
            elif ext == 'json':
                file_type = 'json'
            else:
                return (False, f"Unsupported file extension: .{ext}")
        
        file_type = file_type.lower()
        
        # Check if file exists
        if not self.exists():
            return (False, f"File does not exist: {self}")
        
        if not self.is_file():
            return (False, f"Not a file: {self}")
        
        # Validate JSON (built-in)
        if file_type == 'json':
            try:
                json.loads(self.read_text(encoding='utf-8'))
                return (True, None)
            except json.JSONDecodeError as e:
                return (False, f"Invalid JSON: {e}")
            except Exception as e:
                return (False, f"Error reading file: {e}")
        
        # Validate YAML (requires PyYAML)
        elif file_type in ('yaml', 'yml'):
            if not YAML_AVAILABLE:
                error_msg = "PyYAML library not installed. Install with: pip install pyyaml"
                if strict:
                    raise ImportError(error_msg)
                return (False, error_msg)
            
            try:
                yaml.safe_load(self.read_text(encoding='utf-8'))  # type: ignore
                return (True, None)
            except yaml.YAMLError as e:  # type: ignore
                return (False, f"Invalid YAML: {e}")
            except Exception as e:
                return (False, f"Error reading file: {e}")
        
        # Validate TOML (built-in Python 3.11+ or requires tomli)
        elif file_type == 'toml':
            if not TOML_AVAILABLE:
                if sys.version_info >= (3, 11):
                    error_msg = "TOML support error (this shouldn't happen on Python 3.11+)"
                else:
                    error_msg = "tomli library not installed. Install with: pip install tomli"
                if strict:
                    raise ImportError(error_msg)
                return (False, error_msg)
            
            try:
                if sys.version_info >= (3, 11):
                    import tomllib
                    tomllib.loads(self.read_text(encoding='utf-8'))
                else:
                    import tomli
                    tomli.loads(self.read_text(encoding='utf-8'))
                return (True, None)
            except Exception as e:
                return (False, f"Invalid TOML: {e}")
        
        # Validate INI (built-in)
        elif file_type == 'ini':
            try:
                config = configparser.ConfigParser()  # type: ignore
                config.read_string(self.read_text(encoding='utf-8'))
                return (True, None)
            except configparser.Error as e:  # type: ignore
                return (False, f"Invalid INI: {e}")
            except Exception as e:
                return (False, f"Error reading file: {e}")
        
        else:
            return (False, f"Unsupported file type: {file_type}. Supported: json, yaml, yml, toml, ini")

    # ===============================================================
    # SEARCH & FILTER
    # ===============================================================
    
    def find_files(self, pattern: str = "*") -> List['Path']:
        """
        Find files matching pattern recursively.
        
        Args:
            pattern: Glob pattern
        
        Returns:
            list: List of matching file paths
        
        Example:
            >>> Path('/tmp').find_files('*.txt')
            [Path('/tmp/file1.txt'), Path('/tmp/sub/file2.txt')]
        """
        try:
            return [Path(p) for p in self.rglob(pattern) if p.is_file()]
        except PermissionError:
            return []
    
    def find_dirs(self, pattern: str = "*") -> List['Path']:
        """
        Find directories matching pattern recursively.
        
        Args:
            pattern: Glob pattern
        
        Returns:
            list: List of matching directory paths
        
        Example:
            >>> Path('/tmp').find_dirs('test*')
            [Path('/tmp/test1'), Path('/tmp/sub/test2')]
        """
        try:
            return [Path(p) for p in self.rglob(pattern) if p.is_dir()]
        except PermissionError:
            return []
    
    def walk(self) -> Iterator[Tuple['Path', List[str], List[str]]]:
        """
        Walk directory tree (similar to os.walk).
        
        Yields:
            tuple: (dirpath, dirnames, filenames)
        
        Example:
            >>> for dirpath, dirs, files in Path('/tmp').walk():
            ...     print(f"Directory: {dirpath}")
            ...     print(f"Subdirs: {dirs}")
            ...     print(f"Files: {files}")
        """
        for root, dirs, files in os.walk(self):
            yield (Path(root), dirs, files)
    
    # ===============================================================
    # COMPARISON
    # ===============================================================
    
    def same_content(self, other: Union[str, _PathBase, 'Path'], chunk_size: int = 8192) -> bool:
        """
        Check if two files have the same content.
        
        Args:
            other: Other file path
            chunk_size: Size of chunks for comparison (default: 8192)
        
        Returns:
            bool: True if contents are identical
        
        Example:
            >>> Path('file1.txt').same_content('file2.txt')
            False
        """
        other = Path(other)
        
        if not self.is_file() or not other.is_file():
            return False
        
        # Quick check: different sizes = different content
        if self.size() != other.size():
            return False
        
        # For small files, compare directly
        if self.size() < 1024 * 1024:  # Less than 1MB
            try:
                return self.read_bytes() == other.read_bytes()
            except (PermissionError, OSError):
                return False
        
        # For large files, use hash comparison
        try:
            return self.hash() == other.hash()
        except (PermissionError, OSError):
            return False

    # ===============================================================
    # show music tag, require 'mutagen' package
    # ===============================================================

    def create_table(self) -> Optional[Table]:  # type: ignore
        """
        Create a Rich Table for displaying music tags.
        Returns:
            Table: Rich Table object or None if Rich is not available
        """

        if RICH_AVAILABLE:
            table = Table(title=f"ID3 Tags - {self.basename()}")  # type: ignore
            table.add_column("Tag", style="cyan", no_wrap=True)
            table.add_column("Type", style="magenta")
            table.add_column("Value", style="yellow")
            table.add_column("Size", style="green")
            return table
        return None

    def show_info(self, table: Optional[Table] = None, exts: Optional[List] = ['mp3', 'mp4', 'm4a', 'flac', 'ogg', 'wav', 'wma', 'aac'], no_rich = False) -> None:  # type: ignore
        """
        Show music file tags (requires 'mutagen' package).

        Args:
            table: Rich Table object to populate (if None, a new one is created)
            exts: List of file extensions to consider as music files

        Returns:
            None

        Example:
            >>> Path('song.mp3').show_info()
            >>> Path('music_dir').show_info()
        """

        if not MUTAGEN_AVAILABLE:
            print("mutagen package is not installed. Please install it before.")
            return None
        
        if self.is_file():
            
            if not self.ext().lower() in [i.lower() for i in exts]:  # type: ignore
                return None
            try:
                audio = ID3(self)  # type: ignore
            except Exception:
                return None

            if RICH_AVAILABLE and not no_rich:            
                table = table or self.create_table()  # type: ignore
                if not table:
                    return None
                for tag_key in audio.keys():
                    tag_value = audio[tag_key]
                    tag_type = type(tag_value).__name__
                    
                    # print(f"tag_value: {tag_value}")
                    # print(f"type(tag_value): {type(tag_value)}")
                    # print(f"dir(tag_value): {dir(tag_value)}")

                    # Value format based on tag type
                    if hasattr(tag_value, "text"):
                        # print(f"tag_value.text: {tag_value.text}")
                        # print(f"type(tag_value.text): {type(tag_value.text)}")
                        value = tag_value.text[0] if isinstance(tag_value.text, list) else str(tag_value.text)

                        # print(f"value: {value}")
                        # print(f"type(value): {type(value)}")
                        # print(f"dir(value): {dir(value)}")

                        if hasattr(value, 'text'):
                            # print(f"value.text: {value.text}")
                            # print(f"type(value.text): {type(value.text)}")
                            # print(f"dir(value.text): {dir(value.text)}")
                            value = value.text  # type: ignore

                    elif hasattr(tag_value, "url"):
                        value = tag_value.url
                    elif hasattr(tag_value, "data"):
                        value = f"Binary data ({len(tag_value.data)} bytes)"
                    else:
                        value = " ".join(tag_value) if isinstance(tag_value, (list or tuple)) else str(tag_value)
                    
                    # Trim the value if it is too long
                    if len(value) > 100:
                        value = value[:97] + "..."
                    
                    # Calculate the size
                    size = "N/A"
                    if hasattr(tag_value, "data"):
                        size = f"{len(tag_value.data)} bytes"  # type: ignore
                    elif hasattr(tag_value, "__sizeof__"):
                        try:
                            size = f"{tag_value.__sizeof__()} bytes"
                        except:
                            size = "Unknown"
                    
                    table.add_row(tag_key, tag_type, value, size)

                    # print("="*os.get_terminal_size()[0])
                
                # Show basic file info
                file_size = os.path.getsize(self)
                file_mtime = os.path.getmtime(self)
                mtime_str = datetime.fromtimestamp(file_mtime).strftime('%Y-%m-%d %H:%M:%S')
                
                console.print(table)
                console.print(f"\n[bold]File Info:[/]")
                console.print(f"  Path: {self}")
                console.print(f"  Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)")
                console.print(f"  Modified: {mtime_str}")
                console.print(f"  Number of tags: {len(audio.keys())}")

            else:
                print(f"ID3 Tags - {self.basename()}")
                print("-" * 40)
                for tag_key in audio.keys():
                    tag_value = audio[tag_key]
                    
                    # Value format based on tag type
                    if hasattr(tag_value, "text"):
                        value = str(tag_value.text)
                    elif hasattr(tag_value, "url"):
                        value = tag_value.url
                    elif hasattr(tag_value, "data"):
                        value = f"Binary data ({len(tag_value.data)} bytes)"
                    else:
                        value = str(tag_value)
                    
                    # Trim the value if it is too long
                    if len(value) > 100:
                        value = value[:97] + "..."
                    
                    print(f"{tag_key}: {value}")
                print("-" * 40)
                file_size = os.path.getsize(self)
                file_mtime = os.path.getmtime(self)
                mtime_str = datetime.fromtimestamp(file_mtime).strftime('%Y-%m-%d %H:%M:%S')
                print(f"File Info:")
                print(f"  Path: {self}")
                print(f"  Size: {file_size:,} bytes ({file_size/1024/1024:.2f} MB)")
                print(f"  Modified: {mtime_str}")
                print(f"  Number of tags: {len(audio.keys())}")

        elif self.is_dir():
            for item in self.iterdir():
                if item.is_file() and item.ext().lower() in [i.lower() for i in exts]:  # type: ignore
                    item.show_info(table)

    def music_tag(self, exts: Optional[List] = ['mp3', 'mp4', 'm4a', 'flac', 'ogg', 'wav', 'wma', 'aac']) -> Optional[dict]:
        """
        Get music file tags (requires 'mutagen' package).
        
        Returns:
            dict: Music tags or None if not a music file or mutagen not installed
        
        Example:
            >>> Path('song.mp3').music_tag()
            {'title': 'Song Title', 'artist': 'Artist Name'}
        """
        
        if not MUTAGEN_AVAILABLE:
            print("mutagen package is not installed. Please install it to before.")
            return None
        
        if self.is_file():
            if not self.ext().lower() in [i.lower() for i in exts]:  # type: ignore
                return None
            try:
                audio = MutagenFile(self)  # type: ignore
                if audio is None:
                    return None
                
                tags = {}
                for key in audio.keys():
                    tags[key] = audio[key]
                return tags
            except Exception:
                return None
        elif self.is_dir():
            all_tags = {}
            for item in self.iterdir():
                if item.is_file() and item.ext().lower() in [i.lower() for i in exts]:  # type: ignore
                    tags = item.music_tag()
                    if tags:
                        all_tags[item.basename()] = tags
            return all_tags

    # ================================================================
    # METADATA
    # ================================================================

    def metadata(self, include_basic: bool = True, raw: bool = False) -> dict:
        """
        Extract metadata from various file types.
        
        Supports: Images (JPEG, PNG, etc.), PDF, Audio (MP3, FLAC, etc.), 
                  Video (MP4, MKV, etc.), Office docs (DOCX, XLSX, PPTX)
        
        Args:
            include_basic: Include basic file info (size, dates, permissions)
            raw: Return raw metadata without processing (file-type specific)
        
        Returns:
            dict: Metadata dictionary with file information
        
        Raises:
            ImportError: If required library is not installed
            ValueError: If file type is not supported or file doesn't exist
        
        Example:
            >>> Path('photo.jpg').metadata()
            {
                'file_type': 'image',
                'format': 'JPEG',
                'size': 2048576,
                'size_human': '2.0 MB',
                'width': 1920,
                'height': 1080,
                'created': 1704067200.0,
                'modified': 1704067200.0,
                'exif': {...}
            }
            
            >>> Path('document.pdf').metadata()
            {
                'file_type': 'pdf',
                'pages': 10,
                'author': 'John Doe',
                'title': 'Report',
                ...
            }
        
        Supported formats and required libraries:
            - Images (JPEG, PNG, GIF, BMP, TIFF): Pillow (pip install Pillow)
            - PDF: PyPDF2 (pip install PyPDF2)
            - Audio/Video (MP3, MP4, FLAC, etc.): mutagen (pip install mutagen)
            - Word (DOCX): python-docx (pip install python-docx)
            - Excel (XLSX): openpyxl (pip install openpyxl)
        """
        if not self.exists():
            raise ValueError(f"File does not exist: {self}")
        
        if not self.is_file():
            raise ValueError(f"Not a file: {self}")
        
        # Basic metadata (always included if requested)
        metadata = {}
        
        if include_basic:
            stat_info = self.stat()
            metadata.update({
                'filename': self.name,
                'path': str(self.absolute()),
                'size': stat_info.st_size,
                'size_human': self.size_human(),
                'created': stat_info.st_ctime,
                'modified': stat_info.st_mtime,
                'accessed': stat_info.st_atime,
                'permissions': oct(stat_info.st_mode)[-3:],
            })
        
        ext = self.ext().lower()
        
        # ===== IMAGE FILES =====
        if ext in ('jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp', 'ico'):
            if not PIL_AVAILABLE:
                raise ImportError(
                    "Pillow library not installed. Install with: pip install Pillow"
                )
            
            try:
                with Image.open(self) as img:
                    metadata['file_type'] = 'image'
                    metadata['format'] = img.format
                    metadata['mode'] = img.mode
                    metadata['width'] = img.width
                    metadata['height'] = img.height
                    metadata['resolution'] = f"{img.width}x{img.height}"
                    
                    # Color info
                    if hasattr(img, 'palette'):
                        metadata['has_palette'] = img.palette is not None
                    
                    # EXIF data for JPEG
                    if hasattr(img, '_getexif') and img._getexif():
                        exif_data = {}
                        for tag_id, value in img._getexif().items():
                            tag = TAGS.get(tag_id, tag_id)
                            exif_data[tag] = value
                        
                        if raw:
                            metadata['exif_raw'] = exif_data
                        else:
                            # Extract common EXIF fields
                            metadata['exif'] = {
                                'camera_make': exif_data.get('Make'),
                                'camera_model': exif_data.get('Model'),
                                'datetime': exif_data.get('DateTime'),
                                'orientation': exif_data.get('Orientation'),
                                'flash': exif_data.get('Flash'),
                                'focal_length': exif_data.get('FocalLength'),
                                'f_number': exif_data.get('FNumber'),
                                'exposure_time': exif_data.get('ExposureTime'),
                                'iso': exif_data.get('ISOSpeedRatings'),
                                'gps': exif_data.get('GPSInfo'),
                            }
                            # Remove None values
                            metadata['exif'] = {k: v for k, v in metadata['exif'].items() if v is not None}
                    
                    # Info dict (additional metadata)
                    if hasattr(img, 'info') and img.info:
                        metadata['info'] = img.info if raw else {
                            'dpi': img.info.get('dpi'),
                            'compression': img.info.get('compression'),
                        }
            
            except Exception as e:
                metadata['error'] = f"Failed to read image metadata: {e}"
        
        # ===== PDF FILES =====
        elif ext == 'pdf':
            if not PYPDF2_AVAILABLE:
                raise ImportError(
                    "PyPDF2 library not installed. Install with: pip install PyPDF2"
                )
            
            try:
                with open(self, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    
                    metadata['file_type'] = 'pdf'
                    metadata['pages'] = len(pdf.pages)
                    
                    # PDF metadata
                    if pdf.metadata:
                        pdf_meta = pdf.metadata
                        if raw:
                            metadata['pdf_info'] = dict(pdf_meta)
                        else:
                            metadata['title'] = pdf_meta.get('/Title', '')
                            metadata['author'] = pdf_meta.get('/Author', '')
                            metadata['subject'] = pdf_meta.get('/Subject', '')
                            metadata['creator'] = pdf_meta.get('/Creator', '')
                            metadata['producer'] = pdf_meta.get('/Producer', '')
                            metadata['creation_date'] = pdf_meta.get('/CreationDate', '')
                            metadata['modification_date'] = pdf_meta.get('/ModDate', '')
                    
                    # Check if encrypted
                    metadata['encrypted'] = pdf.is_encrypted
                    
                    # Get text from first page (preview)
                    if len(pdf.pages) > 0 and not raw:
                        try:
                            first_page = pdf.pages[0]
                            text = first_page.extract_text()
                            metadata['preview'] = text[:200] + '...' if len(text) > 200 else text
                        except:
                            pass
            
            except Exception as e:
                metadata['error'] = f"Failed to read PDF metadata: {e}"
        
        # ===== AUDIO/VIDEO FILES =====
        elif ext in ('mp3', 'flac', 'ogg', 'wav', 'm4a', 'wma', 'aac', 
                      'mp4', 'mkv', 'avi', 'mov', 'wmv', 'flv', 'webm'):
            if not MUTAGEN_AVAILABLE:
                raise ImportError(
                    "mutagen library not installed. Install with: pip install mutagen"
                )
            
            try:
                audio = MutagenFile(self)
                
                if audio is None:
                    metadata['error'] = "Unsupported or corrupted media file"
                else:
                    is_video = ext in ('mp4', 'mkv', 'avi', 'mov', 'wmv', 'flv', 'webm')
                    metadata['file_type'] = 'video' if is_video else 'audio'
                    
                    # Audio info
                    if hasattr(audio, 'info'):
                        info = audio.info
                        metadata['length'] = getattr(info, 'length', 0)
                        metadata['length_human'] = f"{int(metadata['length'] // 60)}:{int(metadata['length'] % 60):02d}"
                        metadata['bitrate'] = getattr(info, 'bitrate', 0)
                        metadata['sample_rate'] = getattr(info, 'sample_rate', 0)
                        metadata['channels'] = getattr(info, 'channels', 0)
                    
                    # Tags/Metadata
                    if raw:
                        metadata['tags'] = dict(audio.tags) if audio.tags else {}
                    else:
                        if audio.tags:
                            # Common tags
                            metadata['title'] = audio.tags.get('title', [None])[0] if audio.tags.get('title') else None
                            metadata['artist'] = audio.tags.get('artist', [None])[0] if audio.tags.get('artist') else None
                            metadata['album'] = audio.tags.get('album', [None])[0] if audio.tags.get('album') else None
                            metadata['date'] = audio.tags.get('date', [None])[0] if audio.tags.get('date') else None
                            metadata['genre'] = audio.tags.get('genre', [None])[0] if audio.tags.get('genre') else None
                            metadata['track'] = audio.tags.get('tracknumber', [None])[0] if audio.tags.get('tracknumber') else None
                            
                            # Remove None values
                            metadata = {k: v for k, v in metadata.items() if v is not None}
            
            except Exception as e:
                metadata['error'] = f"Failed to read media metadata: {e}"
        
        # ===== WORD DOCUMENTS (.docx) =====
        elif ext == 'docx':
            if not PYTHON_DOCX_AVAILABLE:
                raise ImportError(
                    "python-docx library not installed. Install with: pip install python-docx"
                )
            
            try:
                doc = docx.Document(self)
                
                metadata['file_type'] = 'document'
                metadata['paragraphs'] = len(doc.paragraphs)
                metadata['tables'] = len(doc.tables)
                
                # Core properties
                props = doc.core_properties
                metadata['title'] = props.title
                metadata['author'] = props.author
                metadata['subject'] = props.subject
                metadata['keywords'] = props.keywords
                metadata['created'] = props.created.timestamp() if props.created else None
                metadata['modified'] = props.modified.timestamp() if props.modified else None
                metadata['last_modified_by'] = props.last_modified_by
                metadata['revision'] = props.revision
                
                # Word count (approximate)
                if not raw:
                    text = '\n'.join([para.text for para in doc.paragraphs])
                    metadata['words'] = len(text.split())
                    metadata['characters'] = len(text)
            
            except Exception as e:
                metadata['error'] = f"Failed to read DOCX metadata: {e}"
        
        # ===== EXCEL FILES (.xlsx) =====
        elif ext in ('xlsx', 'xlsm'):
            if not OPENPYXL_AVAILABLE:
                raise ImportError(
                    "openpyxl library not installed. Install with: pip install openpyxl"
                )
            
            try:
                wb = load_workbook(self, read_only=True, data_only=True)
                
                metadata['file_type'] = 'spreadsheet'
                metadata['sheets'] = len(wb.sheetnames)
                metadata['sheet_names'] = wb.sheetnames
                
                # Properties
                props = wb.properties
                metadata['title'] = props.title
                metadata['creator'] = props.creator
                metadata['subject'] = props.subject
                metadata['keywords'] = props.keywords
                metadata['created'] = props.created.timestamp() if props.created else None
                metadata['modified'] = props.modified.timestamp() if props.modified else None
                metadata['last_modified_by'] = props.lastModifiedBy
                
                # Sheet info
                if not raw:
                    sheet_info = []
                    for sheet_name in wb.sheetnames:
                        sheet = wb[sheet_name]
                        sheet_info.append({
                            'name': sheet_name,
                            'rows': sheet.max_row,
                            'columns': sheet.max_column,
                        })
                    metadata['sheet_info'] = sheet_info
                
                wb.close()
            
            except Exception as e:
                metadata['error'] = f"Failed to read XLSX metadata: {e}"
        
        # ===== TEXT FILES (basic info) =====
        elif ext in ('txt', 'md', 'rst', 'log', 'csv', 'json', 'xml', 'html', 'css', 'js', 'py', 'java', 'c', 'cpp', 'h'):
            metadata['file_type'] = 'text'
            try:
                metadata['lines'] = self.count_lines()
                content = self.read_text(encoding='utf-8')
                metadata['characters'] = len(content)
                metadata['words'] = len(content.split())
            except Exception as e:
                metadata['error'] = f"Failed to read text file: {e}"
        
        # ===== ARCHIVE FILES =====
        elif ext in ('zip', 'tar', 'gz', 'bz2', 'xz', '7z', 'rar'):
            metadata['file_type'] = 'archive'
            metadata['archive_type'] = ext
            
            if ext == 'zip':
                import zipfile
                try:
                    with zipfile.ZipFile(self, 'r') as zf:
                        metadata['files'] = len(zf.namelist())
                        if raw:
                            metadata['file_list'] = zf.namelist()
                        metadata['compressed_size'] = sum(info.compress_size for info in zf.infolist())
                        metadata['uncompressed_size'] = sum(info.file_size for info in zf.infolist())
                except Exception as e:
                    metadata['error'] = f"Failed to read ZIP: {e}"
        
        # ===== UNKNOWN FILE TYPE =====
        else:
            metadata['file_type'] = 'unknown'
            metadata['extension'] = ext
        
        return metadata

    def metadata_simple(self) -> str:
        """
        Get simple human-readable metadata summary.
        
        Returns:
            str: Formatted metadata string
        
        Example:
            >>> print(Path('photo.jpg').metadata_simple())
            File: photo.jpg
            Type: image (JPEG)
            Size: 2.0 MB
            Dimensions: 1920x1080
            Created: 2024-01-01 10:00:00
        """
        try:
            meta = self.metadata(include_basic=True, raw=False)
            
            lines = [f"File: {meta.get('filename', self.name)}"]
            
            # File type specific info
            file_type = meta.get('file_type', 'unknown')
            
            if file_type == 'image':
                lines.append(f"Type: image ({meta.get('format', 'unknown')})")
                lines.append(f"Size: {meta.get('size_human', 'unknown')}")
                lines.append(f"Dimensions: {meta.get('resolution', 'unknown')}")
            
            elif file_type == 'pdf':
                lines.append(f"Type: PDF")
                lines.append(f"Size: {meta.get('size_human', 'unknown')}")
                lines.append(f"Pages: {meta.get('pages', 'unknown')}")
                if meta.get('author'):
                    lines.append(f"Author: {meta.get('author')}")
            
            elif file_type in ('audio', 'video'):
                lines.append(f"Type: {file_type}")
                lines.append(f"Size: {meta.get('size_human', 'unknown')}")
                if meta.get('length_human'):
                    lines.append(f"Duration: {meta.get('length_human')}")
                if meta.get('artist'):
                    lines.append(f"Artist: {meta.get('artist')}")
                if meta.get('title'):
                    lines.append(f"Title: {meta.get('title')}")
            
            else:
                lines.append(f"Type: {file_type}")
                lines.append(f"Size: {meta.get('size_human', 'unknown')}")
            
            # Common info
            if meta.get('modified'):
                from datetime import datetime
                mod_time = datetime.fromtimestamp(meta['modified'])
                lines.append(f"Modified: {mod_time.strftime('%Y-%m-%d %H:%M:%S')}")
            
            return '\n'.join(lines)
        
        except Exception as e:
            return f"Error getting metadata: {e}"
    
    def email_as_attachment(
        self,
        to: Union[str, List[str]],
        subject: str,
        body: str = '',
        config: Optional[EmailConfig] = None,
        from_addr: Optional[str] = None,
        cc: Optional[Union[str, List[str]]] = None,
        bcc: Optional[Union[str, List[str]]] = None,
        body_html: Optional[str] = None,
        attachment_name: Optional[str] = None,
        inline_images: bool = False
    ) -> bool:
        """
        Send this file as email attachment.
        
        Args:
            to: Recipient email(s)
            subject: Email subject
            body: Email body (plain text)
            config: EmailConfig instance with SMTP settings
            from_addr: Sender email (if None, uses config.username)
            cc: CC recipients
            bcc: BCC recipients
            body_html: HTML body (optional, overrides plain body)
            attachment_name: Custom attachment filename (default: file's name)
            inline_images: Embed images inline (for image files only)
        
        Returns:
            bool: True if sent successfully
        
        Raises:
            ImportError: If email libraries not available
            ValueError: If config is not provided or file doesn't exist
            ConnectionError: If SMTP connection fails
        
        Example:
            >>> config = EmailConfig.gmail('me@gmail.com', 'app_password')
            >>> Path('report.pdf').email_as_attachment(
            ...     to='boss@company.com',
            ...     subject='Monthly Report',
            ...     body='Please find attached report.',
            ...     config=config
            ... )
            True
            
            >>> # Multiple recipients
            >>> Path('invoice.pdf').email_as_attachment(
            ...     to=['client@company.com', 'manager@company.com'],
            ...     subject='Invoice #12345',
            ...     body='Invoice attached.',
            ...     cc='accounting@company.com',
            ...     config=config
            ... )
        """
        if not EMAIL_AVAILABLE:
            raise ImportError("Email modules not available")
        
        if config is None:
            raise ValueError(
                "EmailConfig required. Example: "
                "config = EmailConfig.gmail('user@gmail.com', 'password')"
            )
        
        if not self.exists():
            raise ValueError(f"File does not exist: {self}")
        
        if not self.is_file():
            raise ValueError(f"Not a file: {self}")
        
        # Normalize recipients
        to_list = [to] if isinstance(to, str) else to
        cc_list = [cc] if isinstance(cc, str) else (cc or [])
        bcc_list = [bcc] if isinstance(bcc, str) else (bcc or [])
        
        from_addr = from_addr or config.username
        
        # Create message
        msg = MIMEMultipart('alternative' if body_html else 'mixed')
        msg['From'] = from_addr
        msg['To'] = ', '.join(to_list)
        msg['Subject'] = subject
        
        if cc_list:
            msg['Cc'] = ', '.join(cc_list)
        
        # Add body
        if body_html:
            msg.attach(MIMEText(body, 'plain'))
            msg.attach(MIMEText(body_html, 'html'))
        else:
            msg.attach(MIMEText(body, 'plain'))
        
        # Attach file
        attachment_name = attachment_name or self.name
        
        # Handle inline images
        if inline_images and self.ext().lower() in ('jpg', 'jpeg', 'png', 'gif', 'bmp'):
            try:
                with self.open('rb') as f:
                    img_data = f.read()
                    image = MIMEImage(img_data)
                    image.add_header('Content-ID', f'<{attachment_name}>')
                    image.add_header('Content-Disposition', 'inline', filename=attachment_name)
                    msg.attach(image)
            except Exception as e:
                raise ValueError(f"Failed to attach inline image: {e}")
        else:
            # Regular attachment
            try:
                with self.open('rb') as f:
                    # Detect MIME type based on extension
                    ext = self.ext().lower()
                    
                    if ext in ('jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff'):
                        attachment = MIMEImage(f.read())
                    elif ext in ('mp3', 'wav', 'ogg', 'flac'):
                        attachment = MIMEAudio(f.read())
                    elif ext == 'pdf':
                        attachment = MIMEApplication(f.read(), _subtype='pdf')
                    else:
                        # Generic binary attachment
                        part = MIMEBase('application', 'octet-stream')
                        part.set_payload(f.read())
                        encoders.encode_base64(part)
                        attachment = part
                    
                    attachment.add_header(
                        'Content-Disposition',
                        f'attachment; filename="{attachment_name}"'
                    )
                    msg.attach(attachment)
            
            except Exception as e:
                raise ValueError(f"Failed to attach file: {e}")
        
        # Send email
        try:
            if config.use_ssl:
                server = smtplib.SMTP_SSL(
                    config.smtp_server,
                    config.smtp_port,
                    timeout=config.timeout
                )
            else:
                server = smtplib.SMTP(
                    config.smtp_server,
                    config.smtp_port,
                    timeout=config.timeout
                )
                
                if config.use_tls:
                    server.starttls()
            
            if config.username and config.password:
                server.login(config.username, config.password)
            
            # Send to all recipients
            all_recipients = to_list + cc_list + bcc_list
            server.sendmail(from_addr, all_recipients, msg.as_string())
            server.quit()
            
            return True
        
        except smtplib.SMTPAuthenticationError:
            raise ConnectionError(
                "SMTP authentication failed. Check username/password. "
                "For Gmail, use an App Password: https://myaccount.google.com/apppasswords"
            )
        except smtplib.SMTPException as e:
            raise ConnectionError(f"SMTP error: {e}")
        except Exception as e:
            raise ConnectionError(f"Failed to send email: {e}")

    @staticmethod
    def send_email(
        to: Union[str, List[str]],
        subject: str,
        body: str,
        config: EmailConfig,
        from_addr: Optional[str] = None,
        cc: Optional[Union[str, List[str]]] = None,
        bcc: Optional[Union[str, List[str]]] = None,
        body_html: Optional[str] = None,
        attachments: Optional[List[Union[str, 'Path']]] = None
    ) -> bool:
        """
        Send email with optional multiple attachments.
        
        Args:
            to: Recipient email(s)
            subject: Email subject
            body: Email body (plain text)
            config: EmailConfig instance
            from_addr: Sender email
            cc: CC recipients
            bcc: BCC recipients
            body_html: HTML body
            attachments: List of file paths to attach
        
        Returns:
            bool: True if sent successfully
        
        Example:
            >>> config = EmailConfig.gmail('me@gmail.com', 'password')
            >>> Path.send_email(
            ...     to='boss@company.com',
            ...     subject='Weekly Report',
            ...     body='See attached files.',
            ...     config=config,
            ...     attachments=['report.pdf', 'chart.png']
            ... )
        """
        if not EMAIL_AVAILABLE:
            raise ImportError("Email modules not available")
        
        # Normalize recipients
        to_list = [to] if isinstance(to, str) else to
        cc_list = [cc] if isinstance(cc, str) else (cc or [])
        bcc_list = [bcc] if isinstance(bcc, str) else (bcc or [])
        
        from_addr = from_addr or config.username
        
        # Create message
        msg = MIMEMultipart('alternative' if body_html else 'mixed')
        msg['From'] = from_addr
        msg['To'] = ', '.join(to_list)
        msg['Subject'] = subject
        
        if cc_list:
            msg['Cc'] = ', '.join(cc_list)
        
        # Add body
        if body_html:
            msg.attach(MIMEText(body, 'plain'))
            msg.attach(MIMEText(body_html, 'html'))
        else:
            msg.attach(MIMEText(body, 'plain'))
        
        # Add attachments
        if attachments:
            for file_path in attachments:
                file_path = Path(file_path)
                
                if not file_path.exists():
                    raise ValueError(f"Attachment not found: {file_path}")
                
                try:
                    with file_path.open('rb') as f:
                        ext = file_path.ext().lower()
                        
                        if ext in ('jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff'):
                            attachment = MIMEImage(f.read())
                        elif ext in ('mp3', 'wav', 'ogg', 'flac'):
                            attachment = MIMEAudio(f.read())
                        elif ext == 'pdf':
                            attachment = MIMEApplication(f.read(), _subtype='pdf')
                        else:
                            part = MIMEBase('application', 'octet-stream')
                            part.set_payload(f.read())
                            encoders.encode_base64(part)
                            attachment = part
                        
                        attachment.add_header(
                            'Content-Disposition',
                            f'attachment; filename="{file_path.name}"'
                        )
                        msg.attach(attachment)
                
                except Exception as e:
                    raise ValueError(f"Failed to attach {file_path}: {e}")
        
        # Send email
        try:
            if config.use_ssl:
                server = smtplib.SMTP_SSL(
                    config.smtp_server,
                    config.smtp_port,
                    timeout=config.timeout
                )
            else:
                server = smtplib.SMTP(
                    config.smtp_server,
                    config.smtp_port,
                    timeout=config.timeout
                )
                
                if config.use_tls:
                    server.starttls()
            
            if config.username and config.password:
                server.login(config.username, config.password)
            
            all_recipients = to_list + cc_list + bcc_list
            server.sendmail(from_addr, all_recipients, msg.as_string())
            server.quit()
            
            return True
        
        except smtplib.SMTPAuthenticationError:
            raise ConnectionError(
                "SMTP authentication failed. Check username/password. "
                "For Gmail, use an App Password: https://myaccount.google.com/apppasswords"
            )
        except smtplib.SMTPException as e:
            raise ConnectionError(f"SMTP error: {e}")
        except Exception as e:
            raise ConnectionError(f"Failed to send email: {e}")

# ===================================================================
# PurePath3 - Extended PurePath (for path manipulation without I/O)
# ===================================================================
class PurePath3(PurePath):
    """Extended PurePath with additional utility methods (no I/O operations)"""
    
    def ext(self) -> str:
        """Get file extension without the dot"""
        return self.suffix.lstrip('.')
    
    def basename(self) -> str:
        """Get the base name (filename with extension)"""
        return self.name
    
    def base(self) -> str:
        """Get the base name without extension"""
        return self.stem
    
    def dirname(self) -> str:
        """Get the directory name as string"""
        return str(self.parent)
    
    def join(self, *args) -> 'PurePath3':
        """Join path components"""
        result = self
        for arg in args:
            result = result / arg
        return PurePath3(result)
    
    def split_ext(self) -> Tuple[str, str]:
        """Split path into base and extension"""
        return (str(self.with_suffix('')), self.suffix)
    
    def split_path(self) -> List[str]:
        """Split path into list of components"""
        return list(self.parts)
    
    def change_ext(self, new_ext: str) -> 'PurePath3':
        """Change file extension"""
        if not new_ext.startswith('.'):
            new_ext = '.' + new_ext
        return PurePath3(self.with_suffix(new_ext))

def get_version():
    """Get version from __version__.py file"""
    try:
        version_file = Path(__file__).parent / "__version__.py"
        if version_file.is_file():
            with open(version_file, "r") as f:
                for line in f:
                    if line.strip().startswith("version"):
                        parts = line.split("=")
                        if len(parts) == 2:
                            return parts[1].strip().strip('"').strip("'")
    except Exception as e:
        if os.getenv('TRACEBACK') and os.getenv('TRACEBACK') in ['1', 'true', 'True']:
            print(traceback.format_exc())
        else:
            print(f"ERROR: {e}")
    return "3.0.0"

# ===================================================================
# EXPORTS - Re-export everything from pathlib + new classes
# ===================================================================
__all__ = [
    # Original pathlib exports (except Path which we override)
    'PurePath',
    'PosixPath',
    'WindowsPath',
    'PurePosixPath',
    'PureWindowsPath',
    
    # New extended classes
    'Path',        # Our extended Path class
    'PurePath3',

    'YAML_AVAILABLE',
    'TOML_AVAILABLE',
    'INI_AVAILABLE',
    'PIL_AVAILABLE',
    'PYPDF2_AVAILABLE',
    'MUTAGEN_AVAILABLE',
    'PYTHON_DOCX_AVAILABLE',
    'OPENPYXL_AVAILABLE',
    'RICH_AVAILABLE',

    'EmailConfig',
    'EMAIL_AVAILABLE',
]

__version__ = get_version()
__author__ = 'Hadi Cahyadi'
__description__ = 'Extended pathlib with 40+ additional utility methods'
